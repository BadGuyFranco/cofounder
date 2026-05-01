#!/usr/bin/env python3
"""
Build the Pandoc reference.docx for the resume style.

Reads pandoc's default reference.docx (must be pre-generated to /tmp via
`pandoc --print-default-data-file=reference.docx > /tmp/pandoc-default-reference.docx`),
modifies the named styles to match the resume design system (Set B:
Georgia body + Calibri headings, charcoal palette, hairline section
borders), and writes the result to:

    /cofounder/tools/Documentor/local-generator/templates/resume-reference.docx

Run:
    python3 build-resume-reference.py

Design tokens (single source of truth, edit here):
    BODY_FONT      Georgia
    HEADING_FONT   Calibri
    BODY_SIZE      10.5pt
    NAME_SIZE      16pt
    SECTION_SIZE   11pt
    COMPANY_SIZE   11pt
    SUBLABEL_SIZE  10.5pt
    INK            #1F2937 (charcoal, near-black)
    SECONDARY      #4B5563 (medium gray, for future secondary paragraph style)
    PAGE_MARGIN    0.75 in
    LINE_SPACING   1.15
    PARA_SPACING   4pt after, 0pt before
"""
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Design tokens -----------------------------------------------------------
BODY_FONT = "Georgia"
HEADING_FONT = "Calibri"
BODY_SIZE_PT = 10.5
NAME_SIZE_PT = 22                # larger, left-aligned, accent color
SECTION_SIZE_PT = 11
COMPANY_SIZE_PT = 11
SUBLABEL_SIZE_PT = 10.5
INK_HEX = "1F2937"               # charcoal, body text
SECONDARY_HEX = "595959"         # medium gray, secondary info
ACCENT_HEX = "0F4C5C"            # deep teal, name + section rules + hyperlinks
PAGE_MARGIN_IN = 0.75
LINE_SPACING = 1.15
PARA_SPACE_AFTER_PT = 4
SECTION_BORDER_SIZE = 12         # eighths of a point. 12 = 1.5pt rule
NAME_LETTER_SPACING = 20         # twentieths of a point. 20 = 1pt
SECTION_LETTER_SPACING = 12      # 0.6pt
CONTACT_SIZE_PT = 9.5            # smaller, secondary metadata
TAGLINE_SIZE_PT = 10.5
BULLET_INDENT_IN = 0.25          # bullet text hangs at 0.25 in
# ----------------------------------------------------------------------------

REPO_ROOT = Path(__file__).resolve().parent.parent.parent.parent.parent  # /cofounder
TEMPLATE_DIR = Path(__file__).resolve().parent
SRC_DOCX = Path("/tmp/pandoc-default-reference.docx")
OUT_DOCX = TEMPLATE_DIR / "resume-reference.docx"


def ensure_pandoc_default():
    """Generate pandoc's default reference.docx if missing."""
    if not SRC_DOCX.exists():
        print(f"Generating pandoc default at {SRC_DOCX}")
        with open(SRC_DOCX, "wb") as f:
            subprocess.run(
                ["pandoc", "--print-default-data-file=reference.docx"],
                stdout=f,
                check=True,
            )


def hex_color(hex_str):
    return RGBColor.from_string(hex_str)


def set_font(run_or_style_font, name, size_pt=None, bold=None, italic=None, color_hex=None):
    """Apply font properties to a Font object (style.font or run.font)."""
    run_or_style_font.name = name
    # East-asian font fallback so Word respects the western font everywhere
    rPr = run_or_style_font._element  # actually this only works on a run; we handle styles separately
    if size_pt is not None:
        run_or_style_font.size = Pt(size_pt)
    if bold is not None:
        run_or_style_font.bold = bold
    if italic is not None:
        run_or_style_font.italic = italic
    if color_hex is not None:
        run_or_style_font.color.rgb = hex_color(color_hex)


def force_font_in_style(style, font_name):
    """
    Force the style's rFonts to all variants (ascii, hAnsi, cs, eastAsia)
    so Word doesn't substitute when it sees mixed scripts.
    """
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), font_name)


def set_paragraph_spacing(pf, before_pt=None, after_pt=None, line_spacing=None,
                          line_rule="multiple"):
    if before_pt is not None:
        pf.space_before = Pt(before_pt)
    if after_pt is not None:
        pf.space_after = Pt(after_pt)
    if line_spacing is not None:
        pf.line_spacing = line_spacing
        if line_rule == "multiple":
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def add_bottom_border(style, size_pt=4, color_hex="1F2937"):
    """
    Add a hairline bottom border to a paragraph style.
    size is in eighths of a point (4 = 0.5pt).
    """
    pPr = style.element.get_or_add_pPr()
    # Remove existing pBdr if present, then add fresh
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_pt))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_uppercase(style):
    """Add w:caps element so Word renders text in all caps without
    requiring uppercased source text. ATS reads the underlying text."""
    rPr = style.element.get_or_add_rPr()
    caps = rPr.find(qn("w:caps"))
    if caps is None:
        caps = OxmlElement("w:caps")
        caps.set(qn("w:val"), "true")
        rPr.append(caps)


def set_letter_spacing(style, twentieths_of_point):
    """Add w:spacing for character spacing. 10 = 0.5pt."""
    rPr = style.element.get_or_add_rPr()
    sp = rPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        rPr.append(sp)
    sp.set(qn("w:val"), str(twentieths_of_point))


def set_page_margins(doc, inches):
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def main():
    ensure_pandoc_default()
    print(f"Copying base template to {OUT_DOCX}")
    shutil.copy(SRC_DOCX, OUT_DOCX)

    doc = Document(str(OUT_DOCX))

    # Page margins
    set_page_margins(doc, PAGE_MARGIN_IN)
    print(f"  Page margins: {PAGE_MARGIN_IN} in")

    styles = doc.styles

    # ---- Normal (body) -----------------------------------------------------
    normal = styles["Normal"]
    set_font(normal.font, BODY_FONT, size_pt=BODY_SIZE_PT, color_hex=INK_HEX)
    force_font_in_style(normal, BODY_FONT)
    set_paragraph_spacing(normal.paragraph_format, before_pt=0,
                          after_pt=PARA_SPACE_AFTER_PT, line_spacing=LINE_SPACING)
    print(f"  Normal: {BODY_FONT} {BODY_SIZE_PT}pt #{INK_HEX}, line {LINE_SPACING}")

    # ---- Title (markdown # at top of doc; pandoc may use Title for first H1) -
    # We're using # as the candidate name, which pandoc maps to Heading 1.
    # Title style is a fallback. Configure it consistently.
    if "Title" in [s.name for s in styles]:
        title = styles["Title"]
        set_font(title.font, HEADING_FONT, size_pt=NAME_SIZE_PT,
                 bold=True, color_hex=ACCENT_HEX)
        force_font_in_style(title, HEADING_FONT)
        set_letter_spacing(title, NAME_LETTER_SPACING)
        set_paragraph_spacing(title.paragraph_format, before_pt=0,
                              after_pt=2, line_spacing=1.0)
        title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        print(f"  Title: {HEADING_FONT} {NAME_SIZE_PT}pt bold #{ACCENT_HEX}")

    # ---- Heading 1 (the resume's # = Anthony's name) -----------------------
    h1 = styles["Heading 1"]
    set_font(h1.font, HEADING_FONT, size_pt=NAME_SIZE_PT, bold=True,
             color_hex=ACCENT_HEX)
    force_font_in_style(h1, HEADING_FONT)
    set_letter_spacing(h1, NAME_LETTER_SPACING)
    set_paragraph_spacing(h1.paragraph_format, before_pt=0, after_pt=2,
                          line_spacing=1.0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Strip any default heading border or color
    pPr = h1.element.get_or_add_pPr()
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    print(f"  Heading 1 (name): {HEADING_FONT} {NAME_SIZE_PT}pt bold "
          f"#{ACCENT_HEX} left-aligned +1pt tracking")

    # ---- Heading 2 (resume's ## = section headings) ------------------------
    # Section heading text stays charcoal for readability; the rule and
    # the underlying weight/uppercase do the visual lift. The rule is
    # accent teal, 1.5pt (was 0.5pt hairline).
    h2 = styles["Heading 2"]
    set_font(h2.font, HEADING_FONT, size_pt=SECTION_SIZE_PT, bold=True,
             color_hex=INK_HEX)
    force_font_in_style(h2, HEADING_FONT)
    set_paragraph_spacing(h2.paragraph_format, before_pt=14, after_pt=4,
                          line_spacing=1.0)
    set_uppercase(h2)
    set_letter_spacing(h2, SECTION_LETTER_SPACING)
    add_bottom_border(h2, size_pt=SECTION_BORDER_SIZE, color_hex=ACCENT_HEX)
    print(f"  Heading 2 (section): {HEADING_FONT} {SECTION_SIZE_PT}pt bold "
          f"UPPERCASE +0.6pt tracking, 1.5pt #{ACCENT_HEX} bottom rule")

    # ---- Heading 3 (resume's ### = company / role entries) -----------------
    h3 = styles["Heading 3"]
    set_font(h3.font, BODY_FONT, size_pt=COMPANY_SIZE_PT, bold=True,
             color_hex=INK_HEX)
    force_font_in_style(h3, BODY_FONT)
    set_paragraph_spacing(h3.paragraph_format, before_pt=8, after_pt=2,
                          line_spacing=1.15)
    print(f"  Heading 3 (company): {BODY_FONT} {COMPANY_SIZE_PT}pt bold")

    # ---- Heading 4 (resume's #### = sub-labels like 'Results:') ------------
    if "Heading 4" in [s.name for s in styles]:
        h4 = styles["Heading 4"]
        set_font(h4.font, BODY_FONT, size_pt=SUBLABEL_SIZE_PT, bold=True,
                 italic=True, color_hex=INK_HEX)
        force_font_in_style(h4, BODY_FONT)
        set_paragraph_spacing(h4.paragraph_format, before_pt=4, after_pt=1,
                              line_spacing=1.15)
        print(f"  Heading 4 (sub-label): {BODY_FONT} {SUBLABEL_SIZE_PT}pt "
              f"bold italic")

    # ---- List bullet & list number ----------------------------------------
    # Hanging indents: bullet sits at 0in; text starts at BULLET_INDENT_IN;
    # wrapped lines also align at BULLET_INDENT_IN (so prose flows under
    # prose, not under the bullet character).
    for style_name in ("List Bullet", "List Number", "List Paragraph"):
        if style_name in [s.name for s in styles]:
            ls = styles[style_name]
            set_font(ls.font, BODY_FONT, size_pt=BODY_SIZE_PT, color_hex=INK_HEX)
            force_font_in_style(ls, BODY_FONT)
            set_paragraph_spacing(ls.paragraph_format, before_pt=0, after_pt=2,
                                  line_spacing=1.15)
            ls.paragraph_format.left_indent = Inches(BULLET_INDENT_IN)
            ls.paragraph_format.first_line_indent = Inches(-BULLET_INDENT_IN)
    print(f"  List styles: {BODY_FONT} {BODY_SIZE_PT}pt, hanging indent {BULLET_INDENT_IN}in")

    # ---- Custom paragraph styles for header treatment ---------------------
    # Pandoc divs with custom-style="Tagline" / "Contact" map to these.
    # The MD pattern is:
    #     ::: tagline
    #     {italic-gray tagline text}
    #     :::
    #     ::: contact
    #     {small-gray contact line}
    #     :::
    def add_paragraph_style(name, font_name, size_pt, color_hex,
                            italic=False, bold=False,
                            space_before=0, space_after=PARA_SPACE_AFTER_PT,
                            line_spacing=LINE_SPACING):
        """Idempotent add: replaces existing same-named style."""
        from docx.enum.style import WD_STYLE_TYPE
        # python-docx will reuse existing style with same name
        existing = None
        for s in styles:
            if s.name == name:
                existing = s
                break
        if existing is None:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = existing
        style.base_style = styles["Normal"]
        set_font(style.font, font_name, size_pt=size_pt, color_hex=color_hex,
                 italic=italic, bold=bold)
        force_font_in_style(style, font_name)
        set_paragraph_spacing(style.paragraph_format,
                              before_pt=space_before, after_pt=space_after,
                              line_spacing=line_spacing)
        return style

    add_paragraph_style("Tagline", BODY_FONT, TAGLINE_SIZE_PT, SECONDARY_HEX,
                        italic=True, space_before=0, space_after=2,
                        line_spacing=1.1)
    print(f"  Tagline (custom): {BODY_FONT} {TAGLINE_SIZE_PT}pt italic #{SECONDARY_HEX}")

    add_paragraph_style("Contact", BODY_FONT, CONTACT_SIZE_PT, SECONDARY_HEX,
                        italic=False, space_before=0, space_after=8,
                        line_spacing=1.1)
    print(f"  Contact (custom): {BODY_FONT} {CONTACT_SIZE_PT}pt #{SECONDARY_HEX}")

    add_paragraph_style("Role Headline", HEADING_FONT, SECTION_SIZE_PT, INK_HEX,
                        bold=True, space_before=0, space_after=2,
                        line_spacing=1.0)
    print(f"  Role Headline (custom): {HEADING_FONT} {SECTION_SIZE_PT}pt bold "
          f"#{INK_HEX}")

    # ---- Hyperlink color ---------------------------------------------------
    if "Hyperlink" in [s.name for s in styles]:
        hl = styles["Hyperlink"]
        set_font(hl.font, BODY_FONT, color_hex=ACCENT_HEX)
        # Word default underlines hyperlinks; keep that for clarity.
        # Tinted to accent teal so links read as a designed element,
        # not as Windows-default blue.

    doc.save(str(OUT_DOCX))
    size_kb = OUT_DOCX.stat().st_size / 1024
    print(f"\nWrote: {OUT_DOCX}  ({size_kb:.1f} KB)")


if __name__ == "__main__":
    main()
